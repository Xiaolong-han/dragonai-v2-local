"""文件系统工具 - 文件读写、搜索和文档解析"""

import asyncio
import base64
import fnmatch
import logging
from pathlib import Path
from typing import Optional, List

from langchain_core.tools import tool

from app.config import settings
from app.storage.sandbox import FileSandbox

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp"})
IMAGE_MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}

STORAGE_DIR = Path(settings.storage_dir).resolve()


def _resolve_path(file_path: str, for_write: bool = False) -> Path:
    """解析文件路径（使用沙箱验证）
    
    Args:
        file_path: 文件路径
        for_write: 是否用于写入操作
        
    Returns:
        解析后的 Path 对象
        
    Raises:
        ValueError: 路径不合法或在沙箱外
    """
    try:
        if for_write:
            return FileSandbox.validate_path_for_write(file_path)
        return FileSandbox.validate_path(file_path)
    except PermissionError as e:
        raise ValueError(str(e))


def _to_virtual_path(path: Path) -> str:
    """将物理路径转换为虚拟路径"""
    return FileSandbox.to_virtual_path(path)


def _format_with_line_numbers(lines: List[str], start_line: int = 1) -> str:
    """格式化内容并添加行号"""
    max_line_num = start_line + len(lines) - 1
    line_num_width = len(str(max_line_num))
    
    result = []
    for i, line in enumerate(lines):
        line_num = start_line + i
        result.append(f"{line_num:>{line_num_width}}→{line}")
    
    return "\n".join(result)


def _format_size(size: int) -> str:
    """格式化文件大小"""
    for unit in ["B", "KB", "MB", "GB"]:
        if size < 1024:
            return f"{size:.1f}{unit}"
        size /= 1024
    return f"{size:.1f}TB"


def _ls_sync(path: str) -> str:
    """同步列出目录"""
    if path == "/":
        items = _list_dir(STORAGE_DIR)
        return "\n".join(items) if items else "storage 目录为空"
    
    resolved = _resolve_path(path)
    
    if not resolved.exists():
        return f"错误：目录 '{path}' 不存在"
    
    if not resolved.is_dir():
        return f"错误：'{path}' 不是目录"
    
    items = _list_dir(resolved)
    return "\n".join(items) if items else f"目录 '{path}' 为空"


@tool
async def ls(path: str = "/") -> str:
    """列出目录中的文件和子目录。

    用于探索文件系统，查找需要读取或编辑的文件。
    在使用 read_file 或 edit_file 之前，通常应先使用此工具。

    Args:
        path: 目录路径，如 /documents 或 /skills。
              默认为 "/"，显示 storage 目录下的所有内容。

    Returns:
        目录内容列表，包含文件名、类型和大小信息。
    """
    try:
        return await asyncio.to_thread(_ls_sync, path)
    except ValueError as e:
        return f"错误：{e}"
    except Exception as e:
        return f"列出目录时出错: {e}"


def _list_dir(dir_path: Path) -> List[str]:
    """列出目录内容"""
    items = []
    try:
        for child in sorted(dir_path.iterdir()):
            if child.is_dir():
                items.append(f"📁 {child.name}/")
            else:
                size = child.stat().st_size
                size_str = _format_size(size)
                items.append(f"📄 {child.name} ({size_str})")
    except PermissionError:
        items.append("⚠️ 无权限访问")
    return items


@tool
async def read_file(
    file_path: str,
    offset: int = 0,
    limit: int = 100,
) -> str:
    """读取文件内容。

    可以读取文本文件和图片文件。
    
    对于文本文件：
    - 默认读取前 100 行
    - 使用 offset 和 limit 参数分页读取大文件
    - 结果带有行号，从 1 开始
    
    对于图片文件（.png, .jpg, .jpeg, .gif, .webp）：
    - 返回 base64 编码的图片数据
    - Agent 可以直接"看到"图片内容

    Args:
        file_path: 文件路径，如 documents/xxx.txt 或 /documents/xxx.txt。
        offset: 开始读取的行号（0 索引），默认 0。
        limit: 最多读取的行数，默认 100。

    Returns:
        文件内容（带行号）或错误信息。
    """
    try:
        resolved = _resolve_path(file_path)
        
        if not resolved.exists():
            return f"错误：文件 '{file_path}' 不存在"
        
        if not resolved.is_file():
            return f"错误：'{file_path}' 不是文件"
        
        ext = resolved.suffix.lower()
        
        if ext in IMAGE_EXTENSIONS:
            return await _read_image_file(resolved)
        
        return await _read_text_file(resolved, offset, limit)
        
    except ValueError as e:
        return f"错误：{e}"
    except UnicodeDecodeError:
        return f"错误：文件 '{file_path}' 不是有效的文本文件或编码不支持"
    except Exception as e:
        return f"读取文件时出错: {e}"


async def _read_image_file(file_path: Path) -> str:
    """读取图片文件并返回 base64 数据"""
    import aiofiles
    
    mime_type = IMAGE_MEDIA_TYPES.get(file_path.suffix.lower(), "image/png")
    
    async with aiofiles.open(file_path, "rb") as f:
        content = await f.read()
        image_data = base64.b64encode(content).decode("utf-8")
    
    return f"[图片文件: {file_path.name}]\n[Base64 数据: data:{mime_type};base64,{image_data[:100]}...]"


async def _read_text_file(file_path: Path, offset: int, limit: int) -> str:
    """读取文本文件"""
    import aiofiles
    
    async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
        content = await f.read()
    
    if not content.strip():
        return f"文件 '{file_path.name}' 为空"
    
    lines = content.splitlines()
    total_lines = len(lines)
    
    if offset >= total_lines:
        return f"错误：起始行 {offset} 超出文件范围（共 {total_lines} 行）"
    
    end_idx = min(offset + limit, total_lines)
    selected_lines = lines[offset:end_idx]
    
    result = []
    result.append(f"文件: {file_path.name}")
    result.append(f"总行数: {total_lines}")
    result.append(f"显示: 第 {offset + 1} 行 - 第 {end_idx} 行")
    result.append("-" * 50)
    result.append(_format_with_line_numbers(selected_lines, offset + 1))
    
    if end_idx < total_lines:
        result.append("")
        result.append(f"... 还有 {total_lines - end_idx} 行未显示 ...")
    
    return "\n".join(result)


@tool
async def write_file(file_path: str, content: str) -> str:
    """创建新文件并写入内容。

    如果文件已存在，将返回错误。
    如需修改现有文件，请使用 edit_file 工具。

    Args:
        file_path: 文件路径，如 documents/xxx.txt 或 /documents/xxx.txt。
        content: 要写入的文本内容。

    Returns:
        操作结果信息。
    """
    import aiofiles
    
    try:
        resolved = _resolve_path(file_path, for_write=True)
        
        if resolved.exists():
            return f"错误：文件 '{file_path}' 已存在。如需修改，请先读取后使用 edit_file 工具。"
        
        resolved.parent.mkdir(parents=True, exist_ok=True)
        
        async with aiofiles.open(resolved, "w", encoding="utf-8") as f:
            await f.write(content)
        
        logger.info(f"[SANDBOX] File created: {resolved}")
        virtual_path = _to_virtual_path(resolved)
        return f"文件已创建: {virtual_path}"
        
    except ValueError as e:
        return f"错误：{e}"
    except Exception as e:
        return f"写入文件时出错: {e}"


@tool
async def edit_file(
    file_path: str,
    old_string: str,
    new_string: str,
    replace_all: bool = False,
) -> str:
    """编辑文件，替换指定的文本内容。

    必须先使用 read_file 读取文件后才能编辑。
    old_string 必须与文件中的内容完全匹配（包括空格和缩进）。

    Args:
        file_path: 文件路径，如 documents/xxx.txt 或 /documents/xxx.txt。
        old_string: 要查找和替换的文本。必须完全匹配。
        new_string: 替换后的新文本。
        replace_all: 是否替换所有匹配项。默认 False，此时 old_string 必须唯一。

    Returns:
        操作结果信息。
    """
    import aiofiles
    
    try:
        resolved = _resolve_path(file_path)
        
        if not resolved.exists():
            return f"错误：文件 '{file_path}' 不存在"
        
        if not resolved.is_file():
            return f"错误：'{file_path}' 不是文件"
        
        async with aiofiles.open(resolved, "r", encoding="utf-8") as f:
            content = await f.read()
        
        if old_string not in content:
            return f"错误：未找到要替换的内容。请确保 old_string 与文件内容完全匹配。"
        
        count = content.count(old_string)
        
        if count > 1 and not replace_all:
            return f"错误：找到 {count} 处匹配，但 replace_all=False。请提供更具体的 old_string 或设置 replace_all=True。"
        
        if replace_all:
            new_content = content.replace(old_string, new_string)
            occurrences = count
        else:
            new_content = content.replace(old_string, new_string, 1)
            occurrences = 1
        
        async with aiofiles.open(resolved, "w", encoding="utf-8") as f:
            await f.write(new_content)
        
        virtual_path = _to_virtual_path(resolved)
        return f"成功替换 {occurrences} 处内容: {virtual_path}"
        
    except ValueError as e:
        return f"错误：{e}"
    except Exception as e:
        return f"编辑文件时出错: {e}"


def _glob_sync(pattern: str, path: str) -> str:
    """同步执行 glob 搜索"""
    if pattern.startswith("/"):
        pattern = pattern[1:]
    
    base_dir = STORAGE_DIR if path == "/" else _resolve_path(path)
    
    results = []
    for matched in base_dir.rglob(pattern):
        if matched.is_file():
            relative = matched.relative_to(STORAGE_DIR)
            results.append(f"/storage/{relative.as_posix()}")
    
    if not results:
        return f"未找到匹配 '{pattern}' 的文件"
    
    return "\n".join(sorted(results))


@tool
async def glob(pattern: str, path: str = "/") -> str:
    """查找匹配模式的文件。

    支持标准 glob 模式：
    - * 匹配任意字符
    - ** 匹配任意目录
    - ? 匹配单个字符

    示例：
    - **/*.py - 查找所有 Python 文件
    - *.txt - 查找根目录下的所有 txt 文件
    - documents/**/*.md - 查找 documents 目录下的所有 markdown 文件

    Args:
        pattern: glob 匹配模式。
        path: 搜索的起始目录，默认为 storage 根目录。

    Returns:
        匹配的文件路径列表。
    """
    try:
        return await asyncio.to_thread(_glob_sync, pattern, path)
    except ValueError as e:
        return f"错误：{e}"
    except Exception as e:
        return f"搜索文件时出错: {e}"


@tool
async def grep(
    pattern: str,
    path: str = "/",
    glob_pattern: Optional[str] = None,
    output_mode: str = "files_with_matches",
) -> str:
    """在文件中搜索文本内容。

    搜索字面文本（非正则表达式），返回匹配的文件或内容。

    Args:
        pattern: 要搜索的文本内容。
        path: 搜索的起始目录，默认为 storage 根目录。
        glob_pattern: 文件过滤模式，如 *.py 只搜索 Python 文件。
        output_mode: 输出模式：
            - "files_with_matches": 只返回文件路径（默认）
            - "content": 返回匹配的行内容

    Returns:
        搜索结果。
    """
    import aiofiles
    
    try:
        base_dir = STORAGE_DIR if path == "/" else _resolve_path(path)
        
        results = []
        for file_path in base_dir.rglob("*"):
            if not file_path.is_file():
                continue
            
            if glob_pattern:
                if not fnmatch.fnmatch(file_path.name, glob_pattern):
                    continue
            
            try:
                async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                    content = await f.read()
                
                if pattern in content:
                    relative = file_path.relative_to(STORAGE_DIR)
                    virtual_path = f"/storage/{relative.as_posix()}"
                    
                    if output_mode == "files_with_matches":
                        results.append(virtual_path)
                    else:
                        lines = content.splitlines()
                        for i, line in enumerate(lines, 1):
                            if pattern in line:
                                results.append(f"{virtual_path}:{i}: {line.strip()}")
            
            except (UnicodeDecodeError, PermissionError):
                continue
        
        if not results:
            return f"未找到包含 '{pattern}' 的文件"
        
        return "\n".join(results)
        
    except ValueError as e:
        return f"错误：{e}"
    except Exception as e:
        return f"搜索时出错: {e}"


def _read_pdf_sync(file_path: Path, start_page: int, end_page: Optional[int]) -> tuple:
    """同步读取 PDF 文件，返回 (total_pages, content_parts, total_text, needs_ocr)"""
    from pypdf import PdfReader
    
    reader = PdfReader(str(file_path))
    total_pages = len(reader.pages)
    
    actual_end = end_page if end_page else total_pages
    
    content_parts = []
    content_parts.append(f"PDF 文件: {file_path.name}")
    content_parts.append(f"总页数: {total_pages}")
    content_parts.append(f"读取范围: 第 {start_page} 页 - 第 {actual_end} 页")
    content_parts.append("-" * 50)
    
    total_text = ""
    for page_num in range(start_page - 1, actual_end):
        page = reader.pages[page_num]
        text = page.extract_text()
        if text.strip():
            content_parts.append(f"\n=== 第 {page_num + 1} 页 ===\n")
            content_parts.append(text)
            total_text += text
    
    needs_ocr = not total_text.strip()
    return total_pages, content_parts, total_text, needs_ocr


def _render_pdf_page_sync(file_path: Path, page_num: int) -> str:
    """同步渲染 PDF 页面为 base64 图片"""
    import fitz
    import base64
    
    doc = fitz.open(str(file_path))
    page = doc[page_num]
    mat = fitz.Matrix(2.0, 2.0)
    pix = page.get_pixmap(matrix=mat)
    img_data = pix.tobytes("png")
    doc.close()
    
    img_base64 = base64.b64encode(img_data).decode("utf-8")
    return f"data:image/png;base64,{img_base64}"


@tool
async def read_pdf(file_path: str, start_page: int = 1, end_page: Optional[int] = None, use_ocr: bool = False) -> str:
    """读取 PDF 文件内容。

    当用户上传 PDF 文件并需要提取文字内容时使用此工具。
    支持分页读取，避免一次性读取过多内容。
    
    对于扫描版 PDF（图片型 PDF），会自动检测并提示使用 OCR 模式。

    Args:
        file_path: PDF 文件路径，如 documents/xxx.pdf 或 /documents/xxx.pdf。
        start_page: 开始页码，从 1 开始，默认为 1。
        end_page: 结束页码，不指定则读取到最后一页。
        use_ocr: 是否使用 OCR 模式处理扫描版 PDF。默认 False。
                 当 PDF 是扫描图片时，设置为 True 可提取文字。

    Returns:
        PDF 文件的文本内容，包含页码标记。
    """
    try:
        resolved = _resolve_path(file_path)
    except ValueError as e:
        return f"错误：{e}"
    
    if not resolved.exists():
        return f"错误：文件 '{file_path}' 不存在"
    
    if resolved.suffix.lower() != ".pdf":
        return f"错误：文件 '{file_path}' 不是 PDF 格式"
    
    try:
        total_pages, content_parts, total_text, needs_ocr = await asyncio.to_thread(
            _read_pdf_sync, resolved, start_page, end_page
        )
        
        if start_page < 1:
            start_page = 1
        if start_page > total_pages:
            return f"错误：起始页码 {start_page} 超出范围，PDF 共 {total_pages} 页"
        
        if needs_ocr or use_ocr:
            if not use_ocr and needs_ocr:
                content_parts.append("\n⚠️ 检测到这是扫描版 PDF（无文本层）")
                content_parts.append("请设置 use_ocr=True 来使用 OCR 提取文字")
                return "\n".join(content_parts)
            
            if use_ocr:
                return await _ocr_pdf(resolved, start_page, end_page or total_pages)
        
        return "\n".join(content_parts)
        
    except Exception as e:
        return f"读取 PDF 文件时出错: {str(e)}"


async def _ocr_pdf(file_path: Path, start_page: int, end_page: int) -> str:
    """使用 OCR 处理扫描版 PDF"""
    try:
        import fitz
        from app.llm.dashscope_client import get_dashscope_client
        from app.config import settings
        import base64
        
        total_pages = await asyncio.to_thread(lambda: len(fitz.open(str(file_path))))
        
        content_parts = []
        content_parts.append(f"PDF 文件: {file_path.name}")
        content_parts.append(f"总页数: {total_pages}")
        content_parts.append(f"OCR 读取范围: 第 {start_page} 页 - 第 {end_page} 页")
        content_parts.append("-" * 50)
        
        client = get_dashscope_client()
        
        for page_num in range(start_page - 1, min(end_page, total_pages)):
            img_url = await asyncio.to_thread(_render_pdf_page_sync, file_path, page_num)
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"image": img_url},
                        {"text": "请提取这张图片中的所有文字内容，保持原有格式。只输出文字，不要添加任何说明。"}
                    ]
                }
            ]
            
            response = await client.multimodal_call(
                model=settings.model_vision_ocr,
                messages=messages
            )
            text = client.parse_text_response(response)
            
            content_parts.append(f"\n=== 第 {page_num + 1} 页 ===\n")
            content_parts.append(text)
        
        return "\n".join(content_parts)
        
    except ImportError:
        return "错误：OCR 模式需要安装 PyMuPDF 库。请运行: pip install PyMuPDF"
    except Exception as e:
        return f"OCR 处理 PDF 时出错: {str(e)}"


def _read_word_sync(file_path: Path) -> str:
    """同步读取 Word 文档"""
    from docx import Document
    
    doc = Document(str(file_path))
    
    content_parts = []
    content_parts.append(f"Word 文档: {file_path.name}")
    content_parts.append("-" * 50)
    
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name if para.style else "Normal"
            if "Heading" in style_name:
                content_parts.append(f"\n## {para.text}")
            else:
                content_parts.append(para.text)
    
    if doc.tables:
        content_parts.append("\n" + "=" * 50)
        content_parts.append("表格内容:")
        content_parts.append("=" * 50)
        
        for table_idx, table in enumerate(doc.tables, 1):
            content_parts.append(f"\n--- 表格 {table_idx} ---")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    content_parts.append(row_text)
    
    return "\n".join(content_parts)


@tool
async def read_word(file_path: str) -> str:
    """读取 Word 文档(.docx)内容。

    当用户上传 Word 文档并需要提取文字内容时使用此工具。
    支持读取文档标题、段落、表格等内容。

    Args:
        file_path: Word 文件路径，如 documents/xxx.docx 或 /documents/xxx.docx。

    Returns:
        Word 文档的文本内容，包含结构标记。
    """
    try:
        resolved = _resolve_path(file_path)
    except ValueError as e:
        return f"错误：{e}"
    
    if not resolved.exists():
        return f"错误：文件 '{file_path}' 不存在"
    
    if resolved.suffix.lower() not in [".docx", ".doc"]:
        return f"错误：文件 '{file_path}' 不是 Word 文档格式"
    
    if resolved.suffix.lower() == ".doc":
        return "错误：暂不支持旧版 .doc 格式，请转换为 .docx 格式后重试"
    
    try:
        return await asyncio.to_thread(_read_word_sync, resolved)
    except Exception as e:
        return f"读取 Word 文档时出错: {str(e)}"
